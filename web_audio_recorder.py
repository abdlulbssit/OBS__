from flask import Flask, render_template, send_file, jsonify, request, make_response, send_from_directory
from simple_websocket_server import WebSocketServer, WebSocket
import sounddevice as sd
import numpy as np
import wave
import threading
import time
import os
from datetime import datetime
import pyaudio
import json
from openai import OpenAI
from dotenv import load_dotenv
import tempfile
import asyncio
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from docx2pdf import convert
import win32com.client
import pythoncom
import io
from docx.shared import RGBColor
from flask_cors import CORS
from flask_talisman import Talisman
from werkzeug.utils import secure_filename
import logging
from logging.handlers import RotatingFileHandler

# Load environment variables
load_dotenv()
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

app = Flask(__name__, static_folder='templates/static')
CORS(app)
Talisman(app, content_security_policy={
    'default-src': "'self'",
    'script-src': ["'self'", "'unsafe-inline'", "'unsafe-eval'"],
    'style-src': ["'self'", "'unsafe-inline'"],
    'img-src': ["'self'", 'data:', 'blob:'],
    'media-src': ["'self'", 'data:', 'blob:'],
    'connect-src': ["'self'", 'https://api.openai.com']
})

# Configure logging
if not os.path.exists('logs'):
    os.makedirs('logs')

file_handler = RotatingFileHandler('logs/app.log', maxBytes=10240, backupCount=10)
file_handler.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
))
file_handler.setLevel(logging.INFO)
app.logger.addHandler(file_handler)
app.logger.setLevel(logging.INFO)
app.logger.info('Web Audio Recorder startup')

# Ensure recordings directory exists
if not os.path.exists('recordings'):
    os.makedirs('recordings')

class AudioRecorder:
    def __init__(self):
        self.is_recording = False
        self.mic_data = []
        self.sys_data = []
        self.sample_rate = 44100
        self.chunk_size = 1024
        self.clients = set()
        
        # Volume controls
        self.mic_volume = 1.0
        self.sys_volume = 0.3
        
        # Initialize audio devices
        try:
            # Initialize PyAudio
            self.p = pyaudio.PyAudio()
            
            # Get default microphone
            self.devices = sd.query_devices()
            self.default_mic = sd.default.device[0]
            self.mic_name = self.devices[self.default_mic]['name']
            
            # Find WASAPI loopback device
            self.wasapi_device = None
            for i in range(self.p.get_device_count()):
                device_info = self.p.get_device_info_by_index(i)
                if 'Stereo Mix' in device_info['name'] or 'What U Hear' in device_info['name']:
                    self.wasapi_device = i
                    break
                    
        except Exception as e:
            print(f"Error initializing audio devices: {str(e)}")

    def start_recording(self):
        try:
            self.is_recording = True
            self.mic_data = []
            self.sys_data = []
            
            # Start recording threads
            self.mic_thread = threading.Thread(target=self.record_microphone)
            self.mic_thread.daemon = True
            self.mic_thread.start()
            
            if self.wasapi_device is not None:
                self.system_thread = threading.Thread(target=self.record_system_audio)
                self.system_thread.daemon = True
                self.system_thread.start()
            
            # Notify clients that recording has started
            for client in self.clients:
                client.send_json({'event': 'recording_status', 'status': 'started'})
            
            return True
            
        except Exception as e:
            print(f"Failed to start recording: {str(e)}")
            return False
            
    def stop_recording(self):
        try:
            self.is_recording = False
            
            # Wait for threads to finish
            if hasattr(self, 'mic_thread'):
                self.mic_thread.join(timeout=2.0)
            if hasattr(self, 'system_thread'):
                self.system_thread.join(timeout=2.0)
            
            filename = self.save_recording()
            
            if filename:
                # Generate transcript after recording is complete
                transcript = self.transcribe_recording(filename)
                
                # Send transcript to clients
                if transcript:
                    # Generate summary from transcript
                    summary = self.generate_summary(transcript['text'])
                    
                    # Send both transcript and summary to clients
                    for client in self.clients:
                        client.send_json({
                            'event': 'transcription',
                            'text': transcript['text'],
                            'summary': summary,
                            'is_translation': False
                        })
            
            return filename
            
        except Exception as e:
            print(f"Failed to stop recording: {str(e)}")
            return None
    
    def record_microphone(self):
        try:
            def audio_callback(indata, frames, time, status):
                if status and status.input_overflow:
                    print("Input overflow detected - some audio data may have been lost")
                if self.is_recording:
                    if len(indata.shape) == 2:
                        data = indata.copy() * self.mic_volume
                    else:
                        data = (indata.reshape(-1, 2)) * self.mic_volume
                    self.mic_data.append(data)

            with sd.InputStream(
                device=self.default_mic,
                channels=2,
                samplerate=self.sample_rate,
                callback=audio_callback,
                blocksize=self.chunk_size,
                latency='low'
            ):
                while self.is_recording:
                    time.sleep(0.1)
        except Exception as e:
            print(f"Microphone recording failed: {str(e)}")
            self.is_recording = False

    def transcribe_recording(self, filename):
        """Transcribe a complete recording file and translate to English if needed."""
        try:
            with open(filename, 'rb') as f:
                # Use Whisper model for transcription with English translation
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=f,
                    response_format="text",
                    language="en"  # Force English output
                )
                
                if transcript:
                    return {
                        'text': transcript.strip(),
                        'is_translation': True
                    }
                return None
                
        except Exception as e:
            print(f"Transcription error: {str(e)}")
            return None

    def record_system_audio(self):
        try:
            stream = self.p.open(
                format=pyaudio.paFloat32,
                channels=2,
                rate=self.sample_rate,
                input=True,
                input_device_index=self.wasapi_device,
                frames_per_buffer=self.chunk_size
            )
            
            while self.is_recording:
                data = stream.read(self.chunk_size, exception_on_overflow=False)
                if self.is_recording:
                    audio_data = np.frombuffer(data, dtype=np.float32)
                    audio_data = audio_data.reshape(-1, 2) * self.sys_volume
                    self.sys_data.append(audio_data)
                    
            stream.stop_stream()
            stream.close()
            
        except Exception as e:
            print(f"System audio recording failed: {str(e)}")
            self.is_recording = False
            
    def save_recording(self):
        if not (self.mic_data or self.sys_data):
            return None
            
        try:
            if not os.path.exists("recordings"):
                os.makedirs("recordings")
                
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"recordings/recording_{timestamp}.wav"
            
            if self.mic_data:
                mic_data = np.concatenate(self.mic_data, axis=0)
            else:
                mic_data = np.zeros((0, 2), dtype=np.float32)
                
            if self.sys_data:
                sys_data = np.concatenate(self.sys_data, axis=0)
            else:
                sys_data = np.zeros((0, 2), dtype=np.float32)
            
            if len(mic_data) > 0 and len(sys_data) > 0:
                min_length = min(len(mic_data), len(sys_data))
                mic_data = mic_data[:min_length]
                sys_data = sys_data[:min_length]
                
                delay_samples = int(0.05 * self.sample_rate)
                sys_data = np.pad(sys_data, ((delay_samples, 0), (0, 0)))[:-delay_samples]
                
                mixed_data = mic_data + sys_data
            elif len(mic_data) > 0:
                mixed_data = mic_data
            else:
                mixed_data = sys_data
            
            max_val = np.max(np.abs(mixed_data))
            if max_val > 0:
                mixed_data = mixed_data / max_val
            audio_data = np.int16(mixed_data * 32767)
            
            with wave.open(filename, 'wb') as wave_file:
                wave_file.setnchannels(2)
                wave_file.setsampwidth(2)
                wave_file.setframerate(self.sample_rate)
                wave_file.writeframes(audio_data.tobytes())
                
            return filename
        except Exception as e:
            print(f"Failed to save recording: {str(e)}")
            return None
    
    def update_mic_volume(self, value):
        self.mic_volume = float(value)
    
    def update_sys_volume(self, value):
        self.sys_volume = float(value)
    
    def get_device_info(self):
        return {
            "microphone": self.mic_name,
            "system_audio": "Ready" if self.wasapi_device is not None else "Not available (enable Stereo Mix)"
        }

    def generate_summary(self, transcript):
        """Generate a structured meeting analysis report using GPT-3.5."""
        try:
            print("Generating meeting analysis report from transcript...")
            current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": f"""You are a professional meeting analyst that creates structured meeting analysis reports. 
                    Format the report exactly as follows:

                    Meeting Analysis Report
                    Generated on: {current_time}

                    Meeting Duration
                    - Calculate and state the approximate meeting duration

                    Meeting Summary
                    Write a 2-3 sentence overview of the meeting

                    Meeting Agenda
                    List the main topics discussed with bullet points (-)

                    Key Points
                    List the key decisions and important points with bullet points (-)

                    Recommendations
                    List action items and next steps with bullet points (-)

                    Meeting Minutes
                    Number each point chronologically (1., 2., 3., etc.)

                    End with an "Overall" paragraph summarizing the meeting's effectiveness.

                    Important formatting:
                    - Use clear headings without any # symbols
                    - Use '-' for bullet points
                    - Use numbers (1., 2., etc.) for meeting minutes
                    - Keep the format consistent and clean
                    - Justify all text content"""},
                    {"role": "user", "content": f"Please analyze this meeting transcript and create a structured meeting analysis report:\n\n{transcript}"}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            if response and response.choices:
                summary = response.choices[0].message.content.strip()
                print(f"Meeting analysis report generated successfully: {summary[:100]}...")
                return summary
            
            print("No response from OpenAI for summary generation")
            return None
            
        except Exception as e:
            print(f"Summary generation error: {str(e)}")
            return None

# Initialize the recorder instance
recorder = AudioRecorder()

class SimpleWebSocket(WebSocket):
    def handle(self):
        try:
            data = json.loads(self.data)
            if data['event'] == 'recording_started':
                print("Recording started")
            elif data['event'] == 'recording_stopped':
                print("Recording stopped")
        except json.JSONDecodeError:
            pass

    def connected(self):
        print(f"Client {self.address} connected")

    def handle_close(self):
        print(f"Client {self.address} disconnected")

def get_recordings_list():
    recordings_dir = 'recordings'
    if not os.path.exists(recordings_dir):
        os.makedirs(recordings_dir)
    
    recordings = []
    for filename in os.listdir(recordings_dir):
        if filename.endswith('.wav'):
            filepath = os.path.join(recordings_dir, filename)
            recordings.append({
                'filename': filename,
                'date': datetime.fromtimestamp(os.path.getctime(filepath)).strftime('%Y-%m-%d %H:%M:%S'),
                'size': f"{os.path.getsize(filepath) / 1024:.1f} KB"
            })
    return sorted(recordings, key=lambda x: x['date'], reverse=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/static/<path:path>')
def send_static(path):
    return send_from_directory('templates/static', path)

def create_docx_from_transcription(transcription, timestamp, summary=None):
    # Create a new Document
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Audio Transcription & Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp only if no summary (since summary includes its own timestamp)
    if not summary:
        timestamp_para = doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_para.add_run(f'Generated on: {timestamp}')
        timestamp_run.font.size = Pt(10)
    
    # Add a line break
    doc.add_paragraph()
    
    # Add summary section if available
    if summary:
        # Split summary into sections and format
        lines = summary.split('\n')
        timestamp_added = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle timestamp line
            if line.startswith('Generated on:') and not timestamp_added:
                timestamp_para = doc.add_paragraph()
                timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                timestamp_run = timestamp_para.add_run(line)
                timestamp_run.font.size = Pt(10)
                timestamp_added = True
                continue
            
            # Skip duplicate timestamp lines and "Meeting Analysis Report" title
            if line.startswith('Generated on:') or line == 'Meeting Analysis Report':
                continue
                
            # Handle section headers
            if line.startswith(('Meeting Duration', 'Meeting Summary', 'Meeting Agenda', 
                              'Key Points', 'Recommendations', 'Meeting Minutes', 'Overall')):
                heading = doc.add_heading(line, level=2)
                heading.style.font.color.rgb = RGBColor(44, 62, 80)
            # Handle bullet points
            elif line.startswith('-'):
                doc.add_paragraph(line, style='List Bullet')
            # Handle numbered points
            elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                doc.add_paragraph(line, style='List Number')
            # Handle regular text
            else:
                doc.add_paragraph(line)
        
        # Add a page break between summary and transcription
        doc.add_page_break()
    
    # Add the transcription section
    doc.add_heading('Full Transcription', level=1)
    transcription_para = doc.add_paragraph()
    transcription_para.add_run(transcription)
    
    # Save the document to a bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

@app.route('/upload', methods=['POST'])
def upload_audio():
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        if not audio_file:
            return jsonify({'error': 'Empty audio file'}), 400

        # Create a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        temp_filename = temp_file.name
        temp_file.close()

        try:
            # Save the uploaded file
            audio_file.save(temp_filename)

            # Generate transcription
            with open(temp_filename, 'rb') as audio:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio,
                    response_format="text"
                )

            # Generate summary using GPT
            summary = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that creates concise summaries."},
                    {"role": "user", "content": f"Please create a structured summary of this transcription: {transcript}"}
                ]
            ).choices[0].message.content

            return jsonify({
                'transcription': transcript,
                'summary': summary,
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })

        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_filename)
            except Exception as e:
                app.logger.error(f"Error deleting temporary file: {e}")

    except Exception as e:
        app.logger.error(f"Error processing audio: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_audio(filename):
    return send_file(
        os.path.join('recordings', filename),
        as_attachment=True,
        download_name=filename
    )

@app.route('/api/device-info')
def device_info():
    return jsonify(recorder.get_device_info())

@app.route('/api/recordings')
def get_recordings():
    if not os.path.exists("recordings"):
        return jsonify([])
        
    recordings = sorted(
        [f for f in os.listdir("recordings") if f.endswith('.wav')],
        key=lambda x: os.path.getmtime(os.path.join("recordings", x)),
        reverse=True
    )[:5]
    
    recordings_info = []
    for rec in recordings:
        timestamp = datetime.fromtimestamp(
            os.path.getmtime(os.path.join("recordings", rec))
        ).strftime("%Y-%m-%d %H:%M:%S")
        recordings_info.append({
            "filename": rec,
            "timestamp": timestamp
        })
    
    return jsonify(recordings_info)

@app.route('/api/recording/<path:filename>')
def download_recording(filename):
    return send_file(
        os.path.join("recordings", filename),
        as_attachment=True,
        download_name=filename
    )

@app.route('/generate_summary', methods=['POST'])
def generate_summary():
    """Generate a summary from the transcription."""
    try:
        data = request.json
        transcription = data.get('transcription')
        
        if not transcription:
            return jsonify({'error': 'No transcription provided'}), 400
            
        # Generate the summary
        summary = recorder.generate_summary(transcription)
        
        if summary:
            return jsonify({'summary': summary})
        else:
            return jsonify({'error': 'Failed to generate summary'}), 500
            
    except Exception as e:
        print(f"Error generating summary: {str(e)}")
        return jsonify({
            'error': 'Failed to generate summary',
            'details': str(e)
        }), 500

@app.route('/api/recording/<filename>/transcript', methods=['GET'])
def get_transcript(filename):
    try:
        # Construct the full path to the recording
        recording_path = os.path.join('recordings', filename)
        
        if not os.path.exists(recording_path):
            return jsonify({'error': 'Recording not found'}), 404
            
        # Get the transcript using the existing transcription functionality
        transcript_result = recorder.run_transcription(recording_path)
        
        if transcript_result and 'text' in transcript_result:
            return jsonify({'transcript': transcript_result['text']})
        else:
            return jsonify({'error': 'Failed to generate transcript'}), 500
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def create_docx_document(transcript, summary, filename):
    """Create a DOCX document with the transcript and summary."""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Audio Transcription Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add recording info
        doc.add_paragraph(f'Recording: {filename}')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        
        # Add transcript section
        doc.add_heading('Full Transcript', level=1)
        doc.add_paragraph(transcript)
        
        # Save document
        docx_filename = f"{os.path.splitext(filename)[0]}_report.docx"
        docx_path = os.path.join('recordings', docx_filename)
        doc.save(docx_path)
        return docx_filename
    except Exception as e:
        print(f"Error creating DOCX document: {str(e)}")
        return None

@app.route('/api/recording/<filename>/docx')
def download_docx(filename):
    """Download the transcription and summary as a DOCX file."""
    try:
        # Get the transcript
        audio_path = os.path.join('recordings', filename)
        if not os.path.exists(audio_path):
            return jsonify({'error': 'Recording not found'}), 404
            
        with open(audio_path, 'rb') as f:
            transcript = recorder.transcribe_recording(audio_path)
            
        if not transcript:
            return jsonify({'error': 'Failed to generate transcript'}), 500
            
        # Generate summary
        summary = recorder.generate_summary(transcript['text'])
        if not summary:
            return jsonify({'error': 'Failed to generate summary'}), 500
            
        # Create DOCX document
        docx_filename = create_docx_document(transcript['text'], summary, filename)
        if not docx_filename:
            return jsonify({'error': 'Failed to create DOCX document'}), 500
            
        # Send the file
        return send_file(
            os.path.join('recordings', docx_filename),
            as_attachment=True,
            download_name=docx_filename
        )
        
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download-transcription', methods=['POST'])
def download_transcription():
    try:
        data = request.json
        if not data or 'transcription' not in data:
            return jsonify({'error': 'No transcription provided'}), 400

        format_type = data.get('format', 'docx')
        timestamp = data.get('timestamp', datetime.now().strftime("%Y-%m-%d_%H-%M-%S"))
        
        if format_type == 'docx':
            return create_docx_download(data['transcription'], timestamp)
        elif format_type == 'pdf':
            return create_pdf_download(data['transcription'], timestamp)
        else:
            return jsonify({'error': 'Invalid format type'}), 400

    except Exception as e:
        app.logger.error(f"Error creating download: {e}")
        return jsonify({'error': str(e)}), 500

def create_docx_download(text, timestamp):
    doc = Document()
    doc.add_heading('Transcription', 0)
    doc.add_paragraph(f'Generated on: {timestamp}')
    doc.add_paragraph(text)
    
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return send_file(
        docx_file,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'transcription_{timestamp}.docx'
    )

def create_pdf_download(text, timestamp):
    pdf_file = BytesIO()
    c = canvas.Canvas(pdf_file, pagesize=letter)
    
    # Add title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, "Transcription")
    
    # Add timestamp
    c.setFont("Helvetica", 12)
    c.drawString(50, 730, f"Generated on: {timestamp}")
    
    # Add content
    c.setFont("Helvetica", 12)
    text_object = c.beginText(50, 700)
    for line in text.split('\n'):
        text_object.textLine(line)
    c.drawText(text_object)
    
    c.save()
    pdf_file.seek(0)
    
    return send_file(
        pdf_file,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'transcription_{timestamp}.pdf'
    )

@app.route('/api/download/transcript/docx', methods=['POST'])
def download_transcript_docx():
    try:
        # Get content from form data or JSON
        content = request.form.get('content') or request.json.get('content')
        if not content:
            return jsonify({'error': 'No content provided'}), 400

        # Create DOCX
        doc = Document()
        doc.add_heading('Transcript', 0)
        
        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f'Generated on: {timestamp}')
        
        # Add content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to BytesIO
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        response = send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='transcript.docx'
        )
        
        # Add security headers
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        
        return response
        
    except Exception as e:
        print(f"DOCX generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/summary/docx', methods=['POST'])
def download_summary_docx():
    try:
        # Get content from form data or JSON
        content = request.form.get('content') or request.json.get('content')
        if not content:
            return jsonify({'error': 'No content provided'}), 400

        # Create DOCX
        doc = Document()
        
        # Add title
        title = doc.add_heading('Meeting Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content sections
        lines = content.split('\n')
        timestamp_added = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle timestamp line
            if line.startswith('Generated on:') and not timestamp_added:
                timestamp_para = doc.add_paragraph()
                timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                timestamp_run = timestamp_para.add_run(line)
                timestamp_run.font.size = Pt(10)
                timestamp_added = True
                continue
            
            # Skip duplicate timestamp lines
            if line.startswith('Generated on:'):
                continue
                
            # Handle section headers
            if line.startswith(('Meeting Duration', 'Meeting Summary', 'Meeting Agenda', 
                              'Key Points', 'Recommendations', 'Meeting Minutes', 'Overall')):
                heading = doc.add_heading(line, level=1)
                heading.style.font.color.rgb = RGBColor(44, 62, 80)
            # Handle bullet points
            elif line.startswith('-'):
                doc.add_paragraph(line, style='List Bullet')
            # Handle numbered points
            elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                doc.add_paragraph(line, style='List Number')
            # Handle regular text
            else:
                doc.add_paragraph(line)
        
        # Save to BytesIO
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        response = send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='meeting_summary.docx'
        )
        
        # Add security headers
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        
        return response
        
    except Exception as e:
        print(f"DOCX generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/transcript/pdf', methods=['POST'])
def download_transcript_pdf():
    try:
        # Get content from form data or JSON
        content = request.form.get('content') or request.json.get('content')
        if not content:
            return jsonify({'error': 'No content provided'}), 400

        # Create PDF using reportlab
        pdf_buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=12,
            textColor='grey',
            spaceAfter=20
        )
        
        content_style = ParagraphStyle(
            'Content',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            spaceAfter=12
        )
        
        # Build the PDF content
        elements = []
        
        # Add title
        elements.append(Paragraph('Transcript', title_style))
        
        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        elements.append(Paragraph(f'Generated on: {timestamp}', timestamp_style))
        
        # Add content
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                elements.append(Paragraph(para.strip(), content_style))
        
        # Build the PDF
        doc.build(elements)
        
        # Get the value from the BytesIO buffer
        pdf_buffer.seek(0)
        
        response = send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='transcript.pdf'
        )
        
        # Add security headers
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        
        return response
        
    except Exception as e:
        print(f"PDF generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/summary/get-pdf', methods=['POST'])
def download_summary_pdf():
    try:
        # Get content from form data or JSON
        content = request.form.get('content') or request.json.get('content')
        if not content:
            return jsonify({'error': 'No content provided'}), 400

        # Create PDF using reportlab
        pdf_buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=15,
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            fontName='Helvetica-Bold',
            textColor='black',
            alignment=0  # Left alignment for headings
        )
        
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=12,
            textColor='grey',
            spaceAfter=20,
            alignment=0  # Left alignment
        )
        
        content_style = ParagraphStyle(
            'Content',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            spaceAfter=12,
            alignment=4,  # Full justification
            fontName='Helvetica'
        )
        
        bullet_style = ParagraphStyle(
            'Bullet',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            leftIndent=20,
            spaceAfter=8,
            alignment=4,  # Full justification
            fontName='Helvetica'
        )
        
        # Build the PDF content
        elements = []
        
        # Process content sections
        lines = content.split('\n')
        current_section = None
        current_text = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_text:
                    elements.append(Paragraph('\n'.join(current_text), content_style))
                    current_text = []
                continue
            
            # Handle the title and timestamp
            if line == "Meeting Analysis Report":
                elements.append(Paragraph(line, title_style))
            elif line.startswith("Generated on:"):
                elements.append(Paragraph(line, timestamp_style))
            # Handle section headers (now without # symbols)
            elif line in ["Meeting Duration", "Meeting Summary", "Meeting Agenda", 
                         "Key Points", "Recommendations", "Meeting Minutes"]:
                if current_text:
                    elements.append(Paragraph('\n'.join(current_text), content_style))
                    current_text = []
                elements.append(Paragraph(line, heading_style))
            # Handle bullet points
            elif line.startswith('-'):
                if current_text:
                    elements.append(Paragraph('\n'.join(current_text), content_style))
                    current_text = []
                elements.append(Paragraph(line, bullet_style))
            # Handle numbered points
            elif line[0].isdigit() and line[1] == '.':
                if current_text:
                    elements.append(Paragraph('\n'.join(current_text), content_style))
                    current_text = []
                elements.append(Paragraph(line, bullet_style))
            # Handle regular text
            else:
                if line.startswith("Overall"):
                    if current_text:
                        elements.append(Paragraph('\n'.join(current_text), content_style))
                        current_text = []
                current_text.append(line)
        
        # Add any remaining text
        if current_text:
            elements.append(Paragraph('\n'.join(current_text), content_style))
        
        # Build the PDF
        doc.build(elements)
        
        # Get the value from the BytesIO buffer
        pdf_buffer.seek(0)
        
        # Create response
        response = make_response(pdf_buffer.getvalue())
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'attachment; filename=meeting_analysis.pdf'
        return response
        
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return jsonify({'error': str(e)}), 500

def run_websocket_server():
    server = WebSocketServer('localhost', 8765, SimpleWebSocket)
    server.serve_forever()

if __name__ == '__main__':
    # Start WebSocket server in a separate thread
    websocket_thread = threading.Thread(target=run_websocket_server)
    websocket_thread.daemon = True
    websocket_thread.start()
    
    # Start Flask server
    print("Web Audio Recorder is running at http://localhost:5000")
    app.run(host='localhost', port=5000, debug=True) 